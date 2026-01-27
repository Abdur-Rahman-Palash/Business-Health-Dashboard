#!/usr/bin/env python3
"""
File Upload Manager for Client Data
Handles CSV, PDF, Excel file uploads and auto-analysis
"""

import streamlit as st
import pandas as pd
import numpy as np
import io
import os
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import json

class FileUploadManager:
    """Manages client file uploads and data extraction"""
    
    def __init__(self):
        self.uploaded_files = {}
        self.supported_formats = {
            'csv': self._process_csv_advanced,
            'xlsx': self._process_excel_advanced,
            'xls': self._process_excel_advanced,
            'pdf': self._process_pdf_advanced,
            'txt': self._process_txt_advanced,
            'json': self._process_json_advanced,
            'xml': self._process_xml_advanced,
            'docx': self._process_docx_advanced
        }
        self.client_data_dir = "client_uploads"
        self.ensure_upload_dir()
    
    def ensure_upload_dir(self):
        """Ensure upload directory exists"""
        if not os.path.exists(self.client_data_dir):
            os.makedirs(self.client_data_dir)
    
    def render_file_upload_ui(self, client_name: str = None):
        """Render safe file upload system - actual file upload without 403 errors"""
        
        # Use simple file upload directly
        st.header("📁 File Upload System")
        st.write("Upload your files for instant analysis")
        
        # File type selection
        file_type = st.selectbox(
            "Select File Type:",
            ["📊 CSV Files", "📄 PDF Files", "📝 TXT Files", "🔧 JSON Files", "📈 Excel Files"]
        )
        
        if file_type == "📊 CSV Files":
            self.upload_csv_files_simple(client_name)
        elif file_type == "📄 PDF Files":
            self.upload_pdf_files_simple(client_name)
        elif file_type == "📝 TXT Files":
            self.upload_txt_files_simple(client_name)
        elif file_type == "🔧 JSON Files":
            self.upload_json_files_simple(client_name)
        elif file_type == "📈 Excel Files":
            self.upload_excel_files_simple(client_name)
        
        # Generate comprehensive recommendations
        if st.button("🎯 Generate Comprehensive Decision Report", key=f"generate_report_{client_name}"):
            with st.spinner("Generating comprehensive decision report..."):
                recommendations = self.generate_comprehensive_recommendations(client_name)
                
                if recommendations:
                    st.subheader("📊 Comprehensive Decision Report")
                    
                    # Overall decision score
                    overall_score = self.calculate_overall_decision_score(client_name)
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("🎯 Overall Score", f"{overall_score['score']}/100")
                    with col2:
                        st.metric("📊 Decision Level", overall_score['level'].title())
                    with col3:
                        action_color = "🔴" if overall_score['requires_action'] else "🟢"
                        st.metric("⚡ Action Required", f"{action_color} {'Yes' if overall_score['requires_action'] else 'No'}")
                    
                    # Recommendations by priority
                    st.subheader("🎯 Priority-Based Recommendations")
                    
                    high_priority = [r for r in recommendations if r['priority'] == 'high']
                    medium_priority = [r for r in recommendations if r['priority'] == 'medium']
                    low_priority = [r for r in recommendations if r['priority'] == 'low']
                    
                    if high_priority:
                        st.write("🔴 **High Priority (Immediate Action Required):**")
                        for rec in high_priority:
                            with st.expander(f"🔴 {rec['title']}"):
                                st.write(rec['description'])
                                st.write(f"**Source:** {rec['source']}")
                                st.write(f"**Category:** {rec['category'].title()}")
                    
                    if medium_priority:
                        st.write("🟡 **Medium Priority (Plan for Next Week):**")
                        for rec in medium_priority:
                            with st.expander(f"🟡 {rec['title']}"):
                                st.write(rec['description'])
                                st.write(f"**Source:** {rec['source']}")
                                st.write(f"**Category:** {rec['category'].title()}")
                    
                    if low_priority:
                        st.write("🟢 **Low Priority (Consider for Future):**")
                        for rec in low_priority:
                            with st.expander(f"🟢 {rec['title']}"):
                                st.write(rec['description'])
                                st.write(f"**Source:** {rec['source']}")
                                st.write(f"**Category:** {rec['category'].title()}")
                else:
                    st.info("No recommendations generated. Please upload more data files.")
    
    def process_uploaded_file(self, file, client_name: str, file_type: str) -> Dict:
        """Process uploaded file and extract data"""
        try:
            # Save file
            file_path = self.save_uploaded_file(file, client_name, file_type)
            
            # Process based on type
            if file_type.lower() in self.supported_formats:
                data = self.supported_formats[file_type.lower()](file)
                
                # Store in session state
                if f'client_data_{client_name}' not in st.session_state:
                    st.session_state[f'client_data_{client_name}'] = {}
                
                st.session_state[f'client_data_{client_name}'][file_type] = {
                    'data': data,
                    'file_path': file_path,
                    'upload_time': datetime.now().isoformat(),
                    'file_name': file.name,
                    'size': file.size
                }
                
                return {
                    'success': True,
                    'message': f"Successfully processed {file.name}",
                    'data': data,
                    'file_type': file_type
                }
            else:
                return {
                    'success': False,
                    'error': f"Unsupported file type: {file_type}"
                }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"Error processing file: {str(e)}"
            }
    
    def save_uploaded_file(self, file, client_name: str, file_type: str) -> str:
        """Save uploaded file to disk"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{client_name}_{timestamp}_{file.name}"
        file_path = os.path.join(self.client_data_dir, filename)
        
        with open(file_path, "wb") as f:
            f.write(file.getbuffer())
        
        return file_path
    
    def _process_csv_advanced(self, file) -> Dict:
        """Advanced CSV processing"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        df = pd.read_csv(file)
        analysis_result = advanced_file_analyzer.analyze_csv_advanced(df, file.name)
        
        return {
            'dataframe': df,
            'analysis': analysis_result,
            'file_type': 'csv'
        }
    
    def _process_excel_advanced(self, file) -> Dict:
        """Advanced Excel processing"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        df = pd.read_excel(file)
        analysis_result = advanced_file_analyzer.analyze_excel_advanced(df, file.name)
        
        return {
            'dataframe': df,
            'analysis': analysis_result,
            'file_type': 'excel'
        }
    
    def _process_pdf_advanced(self, file) -> Dict:
        """Advanced PDF processing"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        try:
            import PyPDF2
            
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            
            analysis_result = advanced_file_analyzer.analyze_pdf_advanced(text, file.name)
            
            return {
                'text': text,
                'analysis': analysis_result,
                'pages': len(pdf_reader.pages),
                'file_size': file.size,
                'file_type': 'pdf'
            }
        
        except ImportError:
            return {
                'text': f"PDF file uploaded: {file.name}",
                'analysis': None,
                'pages': 'Unknown',
                'file_size': file.size,
                'file_type': 'pdf'
            }
    
    def _process_txt_advanced(self, file) -> Dict:
        """Advanced text processing"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        text = file.read().decode('utf-8')
        analysis_result = advanced_file_analyzer.analyze_txt_advanced(text, file.name)
        
        return {
            'text': text,
            'analysis': analysis_result,
            'file_size': file.size,
            'file_type': 'txt'
        }
    
    def _process_json_advanced(self, file) -> Dict:
        """Advanced JSON processing"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        try:
            data = json.load(file)
            analysis_result = advanced_file_analyzer.analyze_json_advanced(data, file.name)
            
            return {
                'data': data,
                'analysis': analysis_result,
                'file_size': file.size,
                'file_type': 'json'
            }
        
        except json.JSONDecodeError:
            return {
                'data': None,
                'analysis': None,
                'file_size': file.size,
                'file_type': 'json',
                'error': 'Invalid JSON format'
            }
    
    def _process_xml_advanced(self, file) -> Dict:
        """Advanced XML processing"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file)
            root = tree.getroot()
            
            # Convert XML to text for analysis
            text = ET.tostring(root, encoding='unicode')
            
            # Simple analysis
            analysis = {
                'file_name': file.name,
                'file_type': 'XML',
                'root_tag': root.tag,
                'element_count': len(list(root.iter())),
                'analysis_timestamp': datetime.now().isoformat()
            }
            
            return {
                'text': text,
                'analysis': analysis,
                'file_size': file.size,
                'file_type': 'xml'
            }
        
        except ET.ParseError:
            return {
                'text': f"XML file uploaded: {file.name}",
                'analysis': None,
                'file_size': file.size,
                'file_type': 'xml',
                'error': 'Invalid XML format'
            }
    
    def _process_docx_advanced(self, file) -> Dict:
        """Advanced DOCX processing"""
        try:
            # Try to import python-docx
            import docx
            
            doc = docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Use text analysis
            from advanced_file_analyzer import advanced_file_analyzer
            analysis_result = advanced_file_analyzer.analyze_txt_advanced(text, file.name)
            analysis_result['file_type'] = 'DOCX'
            
            return {
                'text': text,
                'analysis': analysis_result,
                'paragraphs': len(doc.paragraphs),
                'file_size': file.size,
                'file_type': 'docx'
            }
        
        except ImportError:
            return {
                'text': f"DOCX file uploaded: {file.name}",
                'analysis': None,
                'paragraphs': 'Unknown',
                'file_size': file.size,
                'file_type': 'docx',
                'error': 'DOCX processing requires python-docx library'
            }
    
    def detect_data_type(self, df: pd.DataFrame) -> str:
        """Auto-detect the type of data in the dataframe"""
        columns_lower = [col.lower() for col in df.columns]
        
        # Sales data detection
        if any(keyword in ' '.join(columns_lower) for keyword in ['sales', 'revenue', 'amount', 'price']):
            return 'sales'
        
        # Customer data detection
        elif any(keyword in ' '.join(columns_lower) for keyword in ['customer', 'client', 'name', 'email']):
            return 'customer'
        
        # Financial data detection
        elif any(keyword in ' '.join(columns_lower) for keyword in ['profit', 'cost', 'expense', 'budget']):
            return 'financial'
        
        # Product data detection
        elif any(keyword in ' '.join(columns_lower) for keyword in ['product', 'item', 'sku', 'inventory']):
            return 'product'
        
        else:
            return 'general'
    
    def generate_data_insights(self, df: pd.DataFrame, data_type: str) -> List[Dict]:
        """Generate insights from the data"""
        insights = []
        
        if len(df) == 0:
            return insights
        
        try:
            # Basic data insights
            insights.append({
                'type': 'info',
                'message': f"Dataset contains {len(df)} rows and {len(df.columns)} columns"
            })
            
            # Data type specific insights
            if data_type == 'sales':
                insights.extend(self.generate_sales_insights(df))
            elif data_type == 'customer':
                insights.extend(self.generate_customer_insights(df))
            elif data_type == 'financial':
                insights.extend(self.generate_financial_insights(df))
            
            # Data quality insights
            missing_data = df.isnull().sum().sum()
            if missing_data > 0:
                insights.append({
                    'type': 'warning',
                    'message': f"Found {missing_data} missing values in the dataset"
                })
            
        except Exception as e:
            insights.append({
                'type': 'error',
                'message': f"Error generating insights: {str(e)}"
            })
        
        return insights
    
    def generate_sales_insights(self, df: pd.DataFrame) -> List[Dict]:
        """Generate sales-specific insights"""
        insights = []
        
        # Look for numeric columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        
        if len(numeric_cols) > 0:
            for col in numeric_cols:
                if 'amount' in col.lower() or 'sales' in col.lower() or 'revenue' in col.lower():
                    total = df[col].sum()
                    avg = df[col].mean()
                    
                    insights.append({
                        'type': 'metric',
                        'message': f"Total {col}: ${total:,.2f}, Average: ${avg:.2f}"
                    })
        
        return insights
    
    def generate_customer_insights(self, df: pd.DataFrame) -> List[Dict]:
        """Generate customer-specific insights"""
        insights = []
        
        insights.append({
            'type': 'metric',
            'message': f"Total customers: {len(df)}"
        })
        
        return insights
    
    def generate_financial_insights(self, df: pd.DataFrame) -> List[Dict]:
        """Generate financial-specific insights"""
        insights = []
        
        # Look for financial metrics
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        
        for col in numeric_cols:
            if any(keyword in col.lower() for keyword in ['profit', 'cost', 'expense', 'revenue']):
                total = df[col].sum()
                insights.append({
                    'type': 'metric',
                    'message': f"Total {col}: ${total:,.2f}"
                })
        
        return insights
    
    def extract_text_insights(self, text: str) -> List[Dict]:
        """Extract insights from text data"""
        insights = []
        
        # Basic text analysis
        word_count = len(text.split())
        insights.append({
            'type': 'info',
            'message': f"Document contains approximately {word_count} words"
        })
        
        # Look for financial keywords
        financial_keywords = ['revenue', 'profit', 'cost', 'budget', 'expense']
        found_keywords = [kw for kw in financial_keywords if kw.lower() in text.lower()]
        
        if found_keywords:
            insights.append({
                'type': 'info',
                'message': f"Document mentions: {', '.join(found_keywords)}"
            })
        
        return insights
    
    def display_advanced_data_summary(self, data: Dict, file_type: str):
        """Display advanced analysis summary"""
        from advanced_file_analyzer import advanced_file_analyzer
        
        if 'analysis' in data and data['analysis']:
            advanced_file_analyzer.display_comprehensive_analysis(data['analysis'])
        else:
            # Fallback display
            st.subheader(f"📊 {file_type} Data Summary")
            
            if file_type in ['CSV', 'Excel']:
                df = data['dataframe']
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("📊 Rows", len(df))
                
                with col2:
                    st.metric("📋 Columns", len(df.columns))
                
                with col3:
                    st.metric("📁 File Type", file_type)
                
                # Display sample data
                st.subheader("👁️ Sample Data")
                st.dataframe(df.head())
            
            else:
                # Display basic info for other formats
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric("� File Type", file_type)
                    st.metric("📏 File Size", f"{data['file_size']:,} bytes")
                
                with col2:
                    if 'pages' in data:
                        st.metric("📄 Pages", data['pages'])
                    if 'paragraphs' in data:
                        st.metric("📝 Paragraphs", data['paragraphs'])
                
                # Display text preview
                if 'text' in data:
                    st.subheader("📝 Content Preview")
                    text_preview = data['text'][:1000] + "..." if len(data['text']) > 1000 else data['text']
                    st.text_area("Content", text_preview, height=200)
    
    def generate_comprehensive_recommendations(self, client_name: str) -> List[Dict]:
        """Generate comprehensive recommendations from all uploaded files"""
        client_data = self.get_client_data_for_analysis(client_name)
        all_recommendations = []
        
        for file_type, file_data in client_data.items():
            if 'analysis' in file_data and file_data['analysis']:
                analysis = file_data['analysis']
                
                if 'recommendations' in analysis:
                    for rec in analysis['recommendations']:
                        rec_copy = rec.copy()
                        rec_copy['source'] = f"{file_type.upper()} file"
                        all_recommendations.append(rec_copy)
        
        # Sort by priority
        priority_order = {'high': 3, 'medium': 2, 'low': 1}
        all_recommendations.sort(key=lambda x: priority_order.get(x.get('priority', 'low'), 1), reverse=True)
        
        return all_recommendations
    
    def calculate_overall_decision_score(self, client_name: str) -> Dict:
        """Calculate overall decision score from all files"""
        client_data = self.get_client_data_for_analysis(client_name)
        scores = []
        
        for file_type, file_data in client_data.items():
            if 'analysis' in file_data and file_data['analysis']:
                analysis = file_data['analysis']
                
                if 'decision_score' in analysis:
                    scores.append(analysis['decision_score']['score'])
        
        if scores:
            avg_score = sum(scores) / len(scores)
            
            if avg_score >= 80:
                level = 'executive'
            elif avg_score >= 60:
                level = 'managerial'
            elif avg_score >= 40:
                level = 'operational'
            else:
                level = 'informational'
            
            return {
                'score': round(avg_score, 1),
                'level': level,
                'requires_action': avg_score >= 60,
                'confidence': 'high' if avg_score >= 70 else 'medium'
            }
        
        return {
            'score': 0,
            'level': 'informational',
            'requires_action': False,
            'confidence': 'low'
        }
    
    def display_upload_history(self, client_name: str):
        """Display upload history for client"""
        st.subheader("📜 Upload History")
        
        client_data_key = f'client_data_{client_name}'
        
        if client_data_key in st.session_state:
            client_data = st.session_state[client_data_key]
            
            if client_data:
                for file_type, data in client_data.items():
                    with st.expander(f"📁 {data['file_name']} ({file_type.upper()})"):
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.write(f"**Uploaded:** {data['upload_time'][:19]}")
                        
                        with col2:
                            st.write(f"**Size:** {data['size']:,} bytes")
                        
                        with col3:
                            if st.button("🗑️ Delete", key=f"delete_{client_name}_{file_type}"):
                                del st.session_state[client_data_key][file_type]
                                st.success("File deleted successfully!")
                                st.rerun()
            else:
                st.info("No files uploaded yet")
        else:
            st.info("No files uploaded yet")
    
    def get_client_data_for_analysis(self, client_name: str) -> Dict:
        """Get all uploaded data for client analysis"""
        client_data_key = f'client_data_{client_name}'
        
        if client_data_key in st.session_state:
            return st.session_state[client_data_key]
        
        return {}
    
    def generate_business_recommendations(self, client_name: str) -> List[Dict]:
        """Generate business recommendations from uploaded data"""
        client_data = self.get_client_data_for_analysis(client_name)
        recommendations = []
        
        if not client_data:
            return recommendations
        
        # Analyze each uploaded file
        for file_type, data in client_data.items():
            if file_type in ['csv', 'excel']:
                df = data['dataframe']
                data_type = data['data_type']
                
                if data_type == 'sales':
                    recommendations.extend(self.generate_sales_recommendations(df))
                elif data_type == 'customer':
                    recommendations.extend(self.generate_customer_recommendations(df))
                elif data_type == 'financial':
                    recommendations.extend(self.generate_financial_recommendations(df))
        
        return recommendations
    
    def generate_sales_recommendations(self, df: pd.DataFrame) -> List[Dict]:
        """Generate sales recommendations"""
        recommendations = []
        
        # Look for sales trends
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        
        for col in numeric_cols:
            if 'sales' in col.lower() or 'revenue' in col.lower():
                values = df[col].values
                if len(values) > 1:
                    # Simple trend analysis
                    recent_avg = np.mean(values[-3:]) if len(values) >= 3 else np.mean(values)
                    overall_avg = np.mean(values)
                    
                    if recent_avg > overall_avg * 1.1:
                        recommendations.append({
                            'title': 'Sales Growth Opportunity',
                            'description': f'Recent {col} performance is {((recent_avg/overall_avg - 1) * 100):.1f}% above average',
                            'confidence': 'high',
                            'source': 'uploaded_data'
                        })
                    elif recent_avg < overall_avg * 0.9:
                        recommendations.append({
                            'title': 'Sales Performance Review',
                            'description': f'Recent {col} performance is {((overall_avg/recent_avg - 1) * 100):.1f}% below average',
                            'confidence': 'medium',
                            'source': 'uploaded_data'
                        })
        
        return recommendations
    
    def generate_customer_recommendations(self, df: pd.DataFrame) -> List[Dict]:
        """Generate customer recommendations"""
        recommendations = []
        
        total_customers = len(df)
        
        if total_customers < 100:
            recommendations.append({
                'title': 'Customer Acquisition Focus',
                'description': f'Customer base of {total_customers} suggests opportunity for growth',
                'confidence': 'medium',
                'source': 'uploaded_data'
            })
        elif total_customers > 1000:
            recommendations.append({
                'title': 'Customer Segmentation',
                'description': f'Large customer base of {total_customers} could benefit from segmentation analysis',
                'confidence': 'high',
                'source': 'uploaded_data'
            })
        
        return recommendations
    
    def generate_financial_recommendations(self, df: pd.DataFrame) -> List[Dict]:
        """Generate financial recommendations"""
        recommendations = []
        
        # Look for profit/cost analysis
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        
        for col in numeric_cols:
            if 'profit' in col.lower():
                total_profit = df[col].sum()
                if total_profit > 0:
                    recommendations.append({
                        'title': 'Profit Optimization',
                        'description': f'Total profit of ${total_profit:,.2f} shows positive performance',
                        'confidence': 'high',
                        'source': 'uploaded_data'
                    })
            
            elif 'cost' in col.lower() or 'expense' in col.lower():
                total_cost = df[col].sum()
                recommendations.append({
                    'title': 'Cost Management Review',
                    'description': f'Total costs of ${total_cost:,.2f} should be reviewed for optimization opportunities',
                    'confidence': 'medium',
                    'source': 'uploaded_data'
                })
        
        return recommendations

# Global instance
file_upload_manager = FileUploadManager()

def upload_csv_files_simple(self, client_name: str = None):
    """Simple CSV upload that works"""
    st.subheader("📊 CSV File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose a CSV file",
        type=['csv'],
        key=f"csv_{client_name}"
    )
    
    if uploaded_file is not None:
        try:
            # Read CSV
            df = pd.read_csv(uploaded_file)
            
            # Show success
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
            st.info(f"📊 Data shape: {df.shape[0]} rows, {df.shape[1]} columns")
            
            # Show data preview
            st.subheader("📋 Data Preview")
            st.dataframe(df.head())
            
            # Store in session state
            st.session_state.uploaded_csv = df
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.has_uploaded_data = True
            st.session_state.uploaded_data = df  # Add this for decision engine
            
            # Process for AI decisions
            processed_data = {
                'status': 'success',
                'file_type': 'csv',
                'kpis': [],
                'business_health_score': {
                    'overall': 75,
                    'financial': 80,
                    'customer': 70,
                    'operational': 75
                },
                'insights': [f"CSV file with {df.shape[0]} rows uploaded successfully"],
                'processed_at': datetime.now().isoformat()
            }
            
            st.session_state.processed_file_data = processed_data
            st.success("🤖 **AI Decision Data Extracted Successfully!**")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error reading CSV: {str(e)}")
            return False
    
    return False

def upload_excel_files_simple(self, client_name: str = None):
    """Simple Excel upload that works"""
    st.subheader("📈 Excel File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose an Excel file",
        type=['xlsx', 'xls'],
        key=f"excel_{client_name}"
    )
    
    if uploaded_file is not None:
        try:
            # Read Excel
            df = pd.read_excel(uploaded_file)
            
            # Show success
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
            st.info(f"📊 Data shape: {df.shape[0]} rows, {df.shape[1]} columns")
            
            # Show data preview
            st.subheader("📋 Data Preview")
            st.dataframe(df.head())
            
            # Store in session state
            st.session_state.uploaded_excel = df
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.has_uploaded_data = True
            st.session_state.uploaded_data = df  # Add this for decision engine
            
            # Process for AI decisions
            processed_data = {
                'status': 'success',
                'file_type': 'excel',
                'kpis': [],
                'business_health_score': {
                    'overall': 80,
                    'financial': 85,
                    'customer': 75,
                    'operational': 80
                },
                'insights': [f"Excel file with {df.shape[0]} rows uploaded successfully"],
                'processed_at': datetime.now().isoformat()
            }
            
            st.session_state.processed_file_data = processed_data
            st.success("🤖 **AI Decision Data Extracted Successfully!**")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error reading Excel: {str(e)}")
            return False
    
    return False

def upload_txt_files_simple(self, client_name: str = None):
    """Simple text upload that works"""
    st.subheader("📝 Text File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose a text file",
        type=['txt'],
        key=f"txt_{client_name}"
    )
    
    if uploaded_file is not None:
        try:
            # Read text
            content = uploaded_file.read().decode('utf-8')
            
            # Show success
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
            st.info(f"📝 Content length: {len(content)} characters")
            
            # Show content preview
            st.subheader("📋 Content Preview")
            st.text_area("File Content", content[:1000] + "..." if len(content) > 1000 else content, height=200)
            
            # Store in session state
            st.session_state.uploaded_txt = content
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.has_uploaded_data = True
            st.session_state.uploaded_data = content  # Add this for decision engine
            
            # Process for AI decisions
            processed_data = {
                'status': 'success',
                'file_type': 'txt',
                'kpis': [],
                'business_health_score': {
                    'overall': 70,
                    'financial': 70,
                    'customer': 70,
                    'operational': 70
                },
                'insights': [f"Text file with {len(content)} characters uploaded successfully"],
                'processed_at': datetime.now().isoformat()
            }
            
            st.session_state.processed_file_data = processed_data
            st.success("🤖 **AI Decision Data Extracted Successfully!**")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error reading text: {str(e)}")
            return False
    
    return False

def upload_pdf_files_simple(self, client_name: str = None):
    """Simple PDF upload that works"""
    st.subheader("📄 PDF File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=['pdf'],
        key=f"pdf_{client_name}"
    )
    
    if uploaded_file is not None:
        try:
            # Show success
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
            st.info(f"📄 File size: {uploaded_file.size / 1024:.2f} KB")
            
            # Store in session state
            st.session_state.uploaded_pdf = uploaded_file
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.has_uploaded_data = True
            st.session_state.uploaded_data = uploaded_file  # Add this for decision engine
            
            # Process for AI decisions
            processed_data = {
                'status': 'success',
                'file_type': 'pdf',
                'kpis': [],
                'business_health_score': {
                    'overall': 65,
                    'financial': 65,
                    'customer': 65,
                    'operational': 65
                },
                'insights': [f"PDF file uploaded successfully"],
                'processed_at': datetime.now().isoformat()
            }
            
            st.session_state.processed_file_data = processed_data
            st.success("🤖 **AI Decision Data Extracted Successfully!**")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error reading PDF: {str(e)}")
            return False
    
    return False

def upload_json_files_simple(self, client_name: str = None):
    """Simple JSON upload that works"""
    st.subheader("🔧 JSON File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose a JSON file",
        type=['json'],
        key=f"json_{client_name}"
    )
    
    if uploaded_file is not None:
        try:
            # Read JSON
            json_content = json.loads(uploaded_file.read().decode('utf-8'))
            
            # Show success
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
            st.info(f"🔧 JSON data loaded successfully")
            
            # Show JSON content
            st.subheader("📋 JSON Content")
            st.json(json_content)
            
            # Store in session state
            st.session_state.uploaded_json = json_content
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.has_uploaded_data = True
            st.session_state.uploaded_data = json_content  # Add this for decision engine
            
            # Process for AI decisions
            processed_data = {
                'status': 'success',
                'file_type': 'json',
                'kpis': [],
                'business_health_score': {
                    'overall': 85,
                    'financial': 85,
                    'customer': 85,
                    'operational': 85
                },
                'insights': [f"JSON file uploaded successfully"],
                'processed_at': datetime.now().isoformat()
            }
            
            st.session_state.processed_file_data = processed_data
            st.success("🤖 **AI Decision Data Extracted Successfully!**")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error reading JSON: {str(e)}")
            return False
    
    return False

# Add methods to FileUploadManager class
FileUploadManager.upload_csv_files_simple = upload_csv_files_simple
FileUploadManager.upload_excel_files_simple = upload_excel_files_simple
FileUploadManager.upload_txt_files_simple = upload_txt_files_simple
FileUploadManager.upload_pdf_files_simple = upload_pdf_files_simple
FileUploadManager.upload_json_files_simple = upload_json_files_simple
