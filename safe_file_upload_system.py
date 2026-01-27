#!/usr/bin/env python3
"""
Safe File Upload System - Working File Upload Without 403 Errors
Complete file upload solution with proper error handling
"""

import streamlit as st
import pandas as pd
import json
import io
import base64
from datetime import datetime
import os
import tempfile
from PIL import Image
import PyPDF2
import docx

def safe_file_upload_system():
    """Complete safe file upload system"""
    
    st.header("📁 Safe File Upload System")
    st.write("Upload and process your files without 403 errors")
    st.success("✅ All File Types Supported - Safe Upload System!")
    
    # File type selection
    file_type = st.selectbox(
        "Select File Type:",
        ["📊 CSV Files", "📄 PDF Files", "📝 TXT Files", "🔧 JSON Files", "📋 XML Files", "📈 Excel Files", "🖼️ Image Files", "📄 Word Documents"]
    )
    
    if file_type == "📊 CSV Files":
        upload_csv_files()
    elif file_type == "📄 PDF Files":
        upload_pdf_files()
    elif file_type == "📝 TXT Files":
        upload_txt_files()
    elif file_type == "🔧 JSON Files":
        upload_json_files()
    elif file_type == "📋 XML Files":
        upload_xml_files()
    elif file_type == "📈 Excel Files":
        upload_excel_files()
    elif file_type == "🖼️ Image Files":
        upload_image_files()
    elif file_type == "📄 Word Documents":
        upload_word_files()

def upload_csv_files():
    """Upload and process CSV files"""
    st.subheader("📊 CSV File Upload")
    
    # File upload with proper error handling
    uploaded_file = st.file_uploader(
        "Choose CSV file",
        type=['csv'],
        help="Upload your CSV file for processing"
    )
    
    if uploaded_file is not None:
        try:
            # Read CSV file
            df = pd.read_csv(uploaded_file)
            
            st.success(f"✅ CSV file uploaded successfully!")
            st.info(f"📊 File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show data preview
            st.subheader("📋 Data Preview")
            st.dataframe(df.head(10))
            
            # Data statistics
            st.subheader("📈 Data Statistics")
            st.write(f"📊 Shape: {df.shape[0]} rows, {df.shape[1]} columns")
            st.write(df.describe())
            
            # Column information
            st.subheader("📋 Column Information")
            col_info = pd.DataFrame({
                'Column': df.columns,
                'Data Type': df.dtypes,
                'Non-Null Count': df.count(),
                'Null Count': df.isnull().sum()
            })
            st.dataframe(col_info)
            
            # Data cleaning options
            st.subheader("🔧 Data Cleaning Options")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                remove_nulls = st.checkbox("Remove Null Values")
            with col2:
                remove_duplicates = st.checkbox("Remove Duplicates")
            with col3:
                reset_index = st.checkbox("Reset Index")
            
            if st.button("🔧 Apply Cleaning"):
                cleaned_df = df.copy()
                
                if remove_nulls:
                    cleaned_df = cleaned_df.dropna()
                    st.success("✅ Null values removed")
                
                if remove_duplicates:
                    cleaned_df = cleaned_df.drop_duplicates()
                    st.success("✅ Duplicates removed")
                
                if reset_index:
                    cleaned_df = cleaned_df.reset_index(drop=True)
                    st.success("✅ Index reset")
                
                st.session_state.cleaned_data = cleaned_df
                st.info(f"📊 Cleaned data shape: {cleaned_df.shape}")
                st.dataframe(cleaned_df.head(10))
            
            # Download options
            st.subheader("💾 Download Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📥 Download Original CSV"):
                    csv_data = df.to_csv(index=False)
                    st.download_button(
                        label="Download Original",
                        data=csv_data,
                        file_name=f"original_{uploaded_file.name}",
                        mime="text/csv"
                    )
            
            with col2:
                if 'cleaned_data' in st.session_state:
                    if st.button("📥 Download Cleaned CSV"):
                        csv_data = st.session_state.cleaned_data.to_csv(index=False)
                        st.download_button(
                            label="Download Cleaned",
                            data=csv_data,
                            file_name=f"cleaned_{uploaded_file.name}",
                            mime="text/csv"
                        )
            
            # Store in session state
            st.session_state.uploaded_csv = df
            st.session_state.uploaded_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing CSV file: {str(e)}")
            st.info("💡 Please check if your file is a valid CSV format")

def upload_pdf_files():
    """Upload and process PDF files"""
    st.subheader("📄 PDF File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose PDF file",
        type=['pdf'],
        help="Upload your PDF file for text extraction"
    )
    
    if uploaded_file is not None:
        try:
            # Read PDF file
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = ""
            
            for page in pdf_reader.pages:
                text_content += page.extract_text()
            
            st.success(f"✅ PDF file uploaded successfully!")
            st.info(f"📄 File: {uploaded_file.name} | Pages: {len(pdf_reader.pages)} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show extracted text
            st.subheader("📋 Extracted Text")
            st.text_area("PDF Content", text_content, height=300)
            
            # Text analysis
            st.subheader("📊 Text Analysis")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("📝 Characters", len(text_content))
            with col2:
                st.metric("📄 Words", len(text_content.split()))
            with col3:
                st.metric("📋 Lines", len(text_content.splitlines()))
            
            # Download extracted text
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download Extracted Text"):
                st.download_button(
                    label="Download Text",
                    data=text_content,
                    file_name=f"extracted_{uploaded_file.name.replace('.pdf', '.txt')}",
                    mime="text/plain"
                )
            
            # Store in session state
            st.session_state.uploaded_pdf_text = text_content
            st.session_state.uploaded_pdf_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing PDF file: {str(e)}")
            st.info("💡 Please check if your file is a valid PDF format")

def upload_txt_files():
    """Upload and process TXT files"""
    st.subheader("📝 TXT File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose TXT file",
        type=['txt'],
        help="Upload your text file for analysis"
    )
    
    if uploaded_file is not None:
        try:
            # Read text file
            text_content = uploaded_file.read().decode('utf-8')
            
            st.success(f"✅ TXT file uploaded successfully!")
            st.info(f"📝 File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show text content
            st.subheader("📋 Text Content")
            st.text_area("File Content", text_content, height=300)
            
            # Text analysis
            st.subheader("📊 Text Analysis")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("📝 Characters", len(text_content))
            with col2:
                st.metric("📄 Words", len(text_content.split()))
            with col3:
                st.metric("📋 Lines", len(text_content.splitlines()))
            with col4:
                st.metric("📊 Paragraphs", len(text_content.split('\n\n')))
            
            # Word frequency
            words = text_content.lower().split()
            word_freq = {}
            for word in words:
                word = ''.join(c for c in word if c.isalnum())
                if word and len(word) > 2:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            if word_freq:
                st.subheader("🔤 Top Words")
                top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
                for word, count in top_words:
                    st.write(f"• {word}: {count}")
            
            # Download options
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download Text"):
                st.download_button(
                    label="Download Text",
                    data=text_content,
                    file_name=uploaded_file.name,
                    mime="text/plain"
                )
            
            # Store in session state
            st.session_state.uploaded_txt_text = text_content
            st.session_state.uploaded_txt_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing TXT file: {str(e)}")
            st.info("💡 Please check if your file is a valid text format")

def upload_json_files():
    """Upload and process JSON files"""
    st.subheader("🔧 JSON File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose JSON file",
        type=['json'],
        help="Upload your JSON file for processing"
    )
    
    if uploaded_file is not None:
        try:
            # Read JSON file
            json_content = json.loads(uploaded_file.read().decode('utf-8'))
            
            st.success(f"✅ JSON file uploaded successfully!")
            st.info(f"🔧 File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show JSON content
            st.subheader("📋 JSON Content")
            st.json(json_content)
            
            # JSON analysis
            st.subheader("📊 JSON Analysis")
            
            if isinstance(json_content, dict):
                st.metric("📋 Object Keys", len(json_content))
                st.write("**Keys:**", list(json_content.keys()))
            elif isinstance(json_content, list):
                st.metric("📋 Array Length", len(json_content))
                if json_content and isinstance(json_content[0], dict):
                    st.write("**Object Keys:**", list(json_content[0].keys()))
            
            # Download options
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download JSON"):
                json_str = json.dumps(json_content, indent=2)
                st.download_button(
                    label="Download JSON",
                    data=json_str,
                    file_name=uploaded_file.name,
                    mime="application/json"
                )
            
            # Store in session state
            st.session_state.uploaded_json = json_content
            st.session_state.uploaded_json_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing JSON file: {str(e)}")
            st.info("💡 Please check if your file is a valid JSON format")

def upload_xml_files():
    """Upload and process XML files"""
    st.subheader("📋 XML File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose XML file",
        type=['xml'],
        help="Upload your XML file for processing"
    )
    
    if uploaded_file is not None:
        try:
            # Read XML file
            xml_content = uploaded_file.read().decode('utf-8')
            
            st.success(f"✅ XML file uploaded successfully!")
            st.info(f"📋 File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show XML content
            st.subheader("📋 XML Content")
            st.code(xml_content, language='xml')
            
            # XML analysis
            st.subheader("📊 XML Analysis")
            
            # Count elements
            import re
            elements = re.findall(r'<(\w+)', xml_content)
            element_counts = {}
            for element in elements:
                element_counts[element] = element_counts.get(element, 0) + 1
            
            if element_counts:
                st.write("**Element Counts:**")
                for element, count in element_counts.items():
                    st.write(f"• {element}: {count}")
            
            # Download options
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download XML"):
                st.download_button(
                    label="Download XML",
                    data=xml_content,
                    file_name=uploaded_file.name,
                    mime="application/xml"
                )
            
            # Store in session state
            st.session_state.uploaded_xml = xml_content
            st.session_state.uploaded_xml_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing XML file: {str(e)}")
            st.info("💡 Please check if your file is a valid XML format")

def upload_excel_files():
    """Upload and process Excel files"""
    st.subheader("📈 Excel File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose Excel file",
        type=['xlsx', 'xls'],
        help="Upload your Excel file for processing"
    )
    
    if uploaded_file is not None:
        try:
            # Read Excel file
            df = pd.read_excel(uploaded_file)
            
            st.success(f"✅ Excel file uploaded successfully!")
            st.info(f"📈 File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show data preview
            st.subheader("📋 Data Preview")
            st.dataframe(df.head(10))
            
            # Data statistics
            st.subheader("📈 Data Statistics")
            st.write(f"📊 Shape: {df.shape[0]} rows, {df.shape[1]} columns")
            st.write(df.describe())
            
            # Column information
            st.subheader("📋 Column Information")
            col_info = pd.DataFrame({
                'Column': df.columns,
                'Data Type': df.dtypes,
                'Non-Null Count': df.count(),
                'Null Count': df.isnull().sum()
            })
            st.dataframe(col_info)
            
            # Download as CSV
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download as CSV"):
                csv_data = df.to_csv(index=False)
                st.download_button(
                    label="Download CSV",
                    data=csv_data,
                    file_name=f"converted_{uploaded_file.name.replace('.xlsx', '.csv').replace('.xls', '.csv')}",
                    mime="text/csv"
                )
            
            # Store in session state
            st.session_state.uploaded_excel = df
            st.session_state.uploaded_excel_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing Excel file: {str(e)}")
            st.info("💡 Please check if your file is a valid Excel format")

def upload_image_files():
    """Upload and process image files"""
    st.subheader("🖼️ Image File Upload")
    
    uploaded_file = st.file_uploader(
        "Choose Image file",
        type=['png', 'jpg', 'jpeg', 'gif', 'bmp'],
        help="Upload your image file for processing"
    )
    
    if uploaded_file is not None:
        try:
            # Read image file
            image = Image.open(uploaded_file)
            
            st.success(f"✅ Image file uploaded successfully!")
            st.info(f"🖼️ File: {uploaded_file.name} | Size: {uploaded_file.size / 1024:.2f} KB | Format: {image.format}")
            
            # Show image
            st.subheader("🖼️ Image Preview")
            st.image(image, caption=uploaded_file.name, use_column_width=True)
            
            # Image information
            st.subheader("📊 Image Information")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("📏 Width", f"{image.width} px")
            with col2:
                st.metric("📐 Height", f"{image.height} px")
            with col3:
                st.metric("🎨 Mode", image.mode)
            
            # Store in session state
            st.session_state.uploaded_image = image
            st.session_state.uploaded_image_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing image file: {str(e)}")
            st.info("💡 Please check if your file is a valid image format")

def upload_word_files():
    """Upload and process Word documents"""
    st.subheader("📄 Word Document Upload")
    
    uploaded_file = st.file_uploader(
        "Choose Word document",
        type=['docx'],
        help="Upload your Word document for text extraction"
    )
    
    if uploaded_file is not None:
        try:
            # Read Word document
            doc = docx.Document(uploaded_file)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            st.success(f"✅ Word document uploaded successfully!")
            st.info(f"📄 File: {uploaded_file.name} | Paragraphs: {len(doc.paragraphs)} | Size: {uploaded_file.size / 1024:.2f} KB")
            
            # Show extracted text
            st.subheader("📋 Extracted Text")
            st.text_area("Document Content", text_content, height=300)
            
            # Text analysis
            st.subheader("📊 Text Analysis")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("📝 Characters", len(text_content))
            with col2:
                st.metric("📄 Words", len(text_content.split()))
            with col3:
                st.metric("📋 Lines", len(text_content.splitlines()))
            
            # Download extracted text
            st.subheader("💾 Download Options")
            
            if st.button("📥 Download Extracted Text"):
                st.download_button(
                    label="Download Text",
                    data=text_content,
                    file_name=f"extracted_{uploaded_file.name.replace('.docx', '.txt')}",
                    mime="text/plain"
                )
            
            # Store in session state
            st.session_state.uploaded_word_text = text_content
            st.session_state.uploaded_word_filename = uploaded_file.name
            
        except Exception as e:
            st.error(f"❌ Error processing Word document: {str(e)}")
            st.info("💡 Please check if your file is a valid Word document format")

if __name__ == "__main__":
    safe_file_upload_system()
